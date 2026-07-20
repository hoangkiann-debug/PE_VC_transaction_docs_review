#!/usr/bin/env python3
"""Prepare isolated blind-evaluation bundles without copying hidden scoring rubrics."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from PIL import Image, ImageDraw


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_MANIFEST = PROJECT_ROOT / "tests" / "evaluation" / "scenarios.json"
SCENARIO_ID_RE = re.compile(r"^[a-z0-9]+(?:-[a-z0-9]+)*$")


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_docx_override(source: Path, output: Path, overrides: dict[str, bytes]) -> None:
    fd, temp_name = tempfile.mkstemp(prefix=".blind-eval-", suffix=".docx", dir=output.parent)
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(source) as archive, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for info in archive.infolist():
                target.writestr(info, overrides.get(info.filename, archive.read(info.filename)))
        temp_path.replace(output)
    finally:
        temp_path.unlink(missing_ok=True)


def replace_paragraph_with_revision(
    paragraph: etree._Element,
    prefix: str,
    deleted: str,
    inserted: str,
    suffix: str,
    revision_id: int,
) -> None:
    properties = paragraph.find(w("pPr"))
    for child in list(paragraph):
        if child is not properties:
            paragraph.remove(child)
    if prefix:
        run = etree.SubElement(paragraph, w("r"))
        etree.SubElement(run, w("t")).text = prefix
    deletion = etree.SubElement(paragraph, w("del"))
    deletion.set(w("id"), str(revision_id))
    deletion.set(w("author"), "Counterparty")
    deletion.set(w("date"), "2026-07-01T00:00:00Z")
    deletion_run = etree.SubElement(deletion, w("r"))
    etree.SubElement(deletion_run, w("delText")).text = deleted
    insertion = etree.SubElement(paragraph, w("ins"))
    insertion.set(w("id"), str(revision_id + 1))
    insertion.set(w("author"), "Counterparty")
    insertion.set(w("date"), "2026-07-01T00:00:01Z")
    insertion_run = etree.SubElement(insertion, w("r"))
    etree.SubElement(insertion_run, w("t")).text = inserted
    if suffix:
        run = etree.SubElement(paragraph, w("r"))
        etree.SubElement(run, w("t")).text = suffix


def generate_native_track_changes(input_dir: Path) -> None:
    prior = input_dir / "prior_clean.docx"
    current = input_dir / "current_markup.docx"
    document = Document()
    document.add_heading("第二轮回购条款", level=1)
    document.add_paragraph("5.2 回购价格为投资款加8%单利。")
    document.add_paragraph("5.3 不设整改期。")
    document.save(prior)
    with zipfile.ZipFile(prior) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    paragraphs = root.xpath(".//w:p", namespaces=NS)
    price = next(paragraph for paragraph in paragraphs if "回购价格" in "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)))
    cure = next(paragraph for paragraph in paragraphs if "整改期" in "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)))
    replace_paragraph_with_revision(price, "5.2 回购价格为投资款加", "8%", "6%", "单利。", 1)
    replace_paragraph_with_revision(cure, "5.3 ", "不设整改期。", "应给予十五个工作日整改期。", "", 3)
    document_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    write_docx_override(prior, current, {"word/document.xml": document_xml})


def generate_scanned_pdf(input_dir: Path) -> None:
    image = Image.new("RGB", (1200, 1600), "white")
    drawing = ImageDraw.Draw(image)
    drawing.text((100, 150), "SCANNED SYNTHETIC INVESTMENT AGREEMENT", fill="black")
    drawing.text((100, 240), "This page contains pixels only and requires OCR.", fill="black")
    image.save(input_dir / "scanned_agreement.pdf", "PDF", resolution=150.0)


def generate_comments_source(input_dir: Path) -> None:
    document = Document()
    document.add_heading("Synthetic Investment Agreement", level=0)
    document.add_heading("Section 8 Reserved Matters", level=1)
    document.add_paragraph("Every expenditure and ordinary-course contract requires the Investor's prior written consent.")
    document.add_paragraph("The Investor's failure to respond is deemed a refusal.")
    document.save(input_dir / "source.docx")


GENERATORS = {
    "native_track_changes": generate_native_track_changes,
    "scanned_pdf": generate_scanned_pdf,
    "comments_source": generate_comments_source,
}


def load_manifest(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def validate_manifest(manifest: dict) -> list[str]:
    errors: list[str] = []
    scenarios = manifest.get("scenarios")
    if manifest.get("schema_version") != 1:
        errors.append("manifest schema_version must be 1")
    if not isinstance(scenarios, list) or len(scenarios) != 12:
        errors.append("manifest must contain exactly 12 scenarios")
        scenarios = []
    identifiers: set[str] = set()
    for scenario in scenarios:
        scenario_id = str(scenario.get("id", ""))
        if not SCENARIO_ID_RE.fullmatch(scenario_id):
            errors.append(f"invalid scenario id: {scenario_id!r}")
        if scenario_id in identifiers:
            errors.append(f"duplicate scenario id: {scenario_id}")
        identifiers.add(scenario_id)
        if not str(scenario.get("user_prompt", "")).strip():
            errors.append(f"{scenario_id}: missing user_prompt")
        if not any(key in scenario for key in ("fixture_dir", "files", "generator")):
            errors.append(f"{scenario_id}: no fixture source")
    return errors


def scenario_map(manifest: dict) -> dict[str, dict]:
    return {str(item["id"]): item for item in manifest.get("scenarios", [])}


def contained_path(root: Path, relative: object, label: str) -> Path:
    root = root.resolve()
    candidate = (root / str(relative)).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"{label} escapes its allowed root: {relative}") from exc
    return candidate


def prepare_scenario(scenario: dict, output_root: Path, project_root: Path = PROJECT_ROOT, replace: bool = False) -> dict[str, object]:
    scenario_id = str(scenario["id"])
    if not SCENARIO_ID_RE.fullmatch(scenario_id):
        raise ValueError(f"invalid scenario id: {scenario_id!r}")
    bundle = output_root / scenario_id
    if bundle.exists():
        if not replace:
            raise FileExistsError(f"bundle already exists: {bundle}")
        shutil.rmtree(bundle)
    input_dir = bundle / "input"
    submission_dir = bundle / "submission"
    input_dir.mkdir(parents=True)
    submission_dir.mkdir()
    if scenario.get("fixture_dir"):
        source = contained_path(project_root, scenario["fixture_dir"], "fixture_dir")
        if not source.is_dir():
            raise FileNotFoundError(f"fixture directory not found: {source}")
        shutil.copytree(source, input_dir, dirs_exist_ok=True)
    for relative, content in scenario.get("files", {}).items():
        target = contained_path(input_dir, relative, "fixture file")
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(str(content), encoding="utf-8")
    generator = scenario.get("generator")
    if generator:
        generator_type = str(generator.get("type", ""))
        if generator_type not in GENERATORS:
            raise ValueError(f"unknown generator: {generator_type}")
        GENERATORS[generator_type](input_dir)
    request_path = bundle / "USER_REQUEST.md"
    request_path.write_text(str(scenario["user_prompt"]).strip() + "\n", encoding="utf-8")
    input_files = sorted(path for path in input_dir.rglob("*") if path.is_file())
    run_manifest = {
        "schema_version": 1,
        "scenario_id": scenario_id,
        "family": scenario.get("family"),
        "language": scenario.get("language"),
        "request_file": "USER_REQUEST.md",
        "input_dir": "input",
        "submission_dir": "submission",
        "input_files": [
            {
                "path": str(path.relative_to(bundle)),
                "sha256": sha256_file(path),
                "size": path.stat().st_size,
            }
            for path in input_files
        ],
        "rubric_included": False,
    }
    (bundle / "run-manifest.json").write_text(json.dumps(run_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"scenario_id": scenario_id, "bundle": str(bundle), "input_file_count": len(input_files)}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--scenario", action="append")
    group.add_argument("--all", action="store_true")
    parser.add_argument("--output-root", type=Path, required=True)
    parser.add_argument("--replace", action="store_true")
    args = parser.parse_args()
    try:
        manifest = load_manifest(args.manifest)
        manifest_errors = validate_manifest(manifest)
        if manifest_errors:
            raise ValueError(f"invalid evaluation manifest: {manifest_errors}")
        scenarios = scenario_map(manifest)
        selected = sorted(scenarios) if args.all else args.scenario
        unknown = [scenario_id for scenario_id in selected if scenario_id not in scenarios]
        if unknown:
            raise ValueError(f"unknown scenario IDs: {unknown}")
        results = [prepare_scenario(scenarios[scenario_id], args.output_root, replace=args.replace) for scenario_id in selected]
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"Could not prepare blind evaluation: {exc}", file=sys.stderr)
        return 2
    print(json.dumps({"prepared": results, "rubric_included": False}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
